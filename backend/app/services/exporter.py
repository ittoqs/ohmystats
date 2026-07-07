from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

def add_table_borders(table):
    # Standard APA style: top and bottom borders for the header row, bottom border for the last row
    tbl = table._tbl

    # Define a helper function to set borders for a row
    def set_row_border(row, top=False, bottom=False):
        tc_list = row.xpath('.//w:tc')
        for tc in tc_list:
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            if top:
                top_border = OxmlElement('w:top')
                top_border.set(qn('w:val'), 'single')
                top_border.set(qn('w:sz'), '4')
                top_border.set(qn('w:space'), '0')
                top_border.set(qn('w:color'), '000000')
                tcBorders.append(top_border)

            if bottom:
                bottom_border = OxmlElement('w:bottom')
                bottom_border.set(qn('w:val'), 'single')
                bottom_border.set(qn('w:sz'), '4')
                bottom_border.set(qn('w:space'), '0')
                bottom_border.set(qn('w:color'), '000000')
                tcBorders.append(bottom_border)

            tcPr.append(tcBorders)

    # Remove all default borders
    for row in tbl.xpath('.//w:tr'):
        for tc in row.xpath('.//w:tc'):
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is not None:
                tcPr.remove(tcBorders)

    # Set specific borders
    rows = tbl.xpath('.//w:tr')
    if len(rows) > 0:
        set_row_border(rows[0], top=True, bottom=True)  # Header row
    if len(rows) > 1:
        set_row_border(rows[-1], bottom=True)           # Last row

def generate_report_docx(analysis_result: dict) -> BytesIO:
    doc = Document()

    # Title
    title = doc.add_heading('Laporan Analisis Regresi ohmystats', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info
    doc.add_paragraph(f"Jenis Regresi: {analysis_result.get('regression_type', 'N/A').replace('_', ' ').title()}")

    metrics = analysis_result.get('metrics', {})
    if metrics:
        doc.add_heading('Ringkasan Model', level=1)
        p = doc.add_paragraph()
        for k, v in metrics.items():
            if v is not None:
                p.add_run(f"{k.replace('_', ' ').title()}: {v:.4f}\n")

    coefficients = analysis_result.get('coefficients', [])
    if coefficients:
        doc.add_heading('Tabel Koefisien', level=1)

        # Determine columns based on available keys
        keys = list(coefficients[0].keys())
        table = doc.add_table(rows=1, cols=len(keys))
        table.style = 'Normal Table'

        # Header
        hdr_cells = table.rows[0].cells
        for i, key in enumerate(keys):
            hdr_cells[i].text = key.replace('_', ' ').title()

        # Data
        for row_data in coefficients:
            row_cells = table.add_row().cells
            for i, key in enumerate(keys):
                val = row_data.get(key, "")
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.4f}"
                else:
                    row_cells[i].text = str(val)

        add_table_borders(table)

    insights = analysis_result.get('insights', '')
    if insights:
        doc.add_heading('Kesimpulan', level=1)
        doc.add_paragraph(insights)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
